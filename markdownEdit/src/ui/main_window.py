from pathlib import Path

from PyQt5.QtCore import QSettings, QTimer
from PyQt5.QtGui import QDragEnterEvent, QDropEvent, QIcon, QKeySequence
from PyQt5.QtPrintSupport import QPrinter
from PyQt5.QtWebEngineWidgets import QWebEnginePage
from PyQt5.QtWidgets import (
    QAction,
    QActionGroup,
    QFileDialog,
    QMainWindow,
    QMessageBox,
    QStatusBar,
    QTabWidget,
)

from src import APP_NAME, __version__
from src.render.md_renderer import MarkdownRenderer
from src.translate.hover import HoverTranslator
from src.translate.worker import TranslationWorker, set_active_worker
from src.ui.editor_tab import EditorTab, VIEW_PREVIEW, VIEW_TRANSLATION
from src.utils.chrome import open_html_in_chrome_translate
from src.utils.paths import dict_path, model_path

from docx import Document


RECENT_MAX = 8


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        icon_path = str(Path(__file__).parent.parent.parent / "resources" / "app_icon.ico")
        self.setWindowIcon(QIcon(icon_path))
        self.setWindowTitle(f"{APP_NAME} v{__version__}")
        self.resize(1200, 760)

        self._settings = QSettings("MarkdownEdit", "App")
        self._renderer = MarkdownRenderer()

        self._dict_path = dict_path()

        self._refresh_timer = QTimer(self)
        self._refresh_timer.setSingleShot(True)
        self._refresh_timer.timeout.connect(self._do_refresh_translation)
        self._refresh_pending = False

        self.tabs = QTabWidget(self)
        self.tabs.setTabsClosable(True)
        self.tabs.setMovable(True)
        self.tabs.setDocumentMode(True)
        self.tabs.tabCloseRequested.connect(self._close_tab)
        self.tabs.currentChanged.connect(self._on_current_tab_changed)
        self.setCentralWidget(self.tabs)

        self.setStatusBar(QStatusBar(self))
        self.setAcceptDrops(True)
        self._build_menus()

        self._add_new_tab()

        self._init_worker()

    # ----- tab management -----

    def _add_new_tab(self, text: str = "", path: Path | None = None) -> EditorTab:
        tab = EditorTab(self._renderer, self)
        tab.set_markdown_text(text)
        tab.current_path = path
        tab.dirtyChanged.connect(self._on_tab_dirty_changed)
        tab.fileDropped.connect(self._on_file_dropped)
        tab.openInChromeRequested.connect(self._open_preview_in_chrome)
        tab.saveHtmlRequested.connect(self._save_preview_as_html)
        tab.savePdfRequested.connect(self._save_preview_as_pdf)
        tab.saveWordRequested.connect(self._save_preview_as_word)
        if self._dict_path.exists():
            tab._hover = HoverTranslator(tab.editor, self._dict_path)
        index = self.tabs.addTab(tab, tab.tab_title)
        self.tabs.setCurrentIndex(index)

        if path:
            self._push_recent(str(path))
        return tab

    def _current_tab(self) -> EditorTab | None:
        idx = self.tabs.currentIndex()
        if idx >= 0:
            w = self.tabs.widget(idx)
            if isinstance(w, EditorTab):
                return w
        return None

    def _close_tab(self, index: int) -> None:
        tab = self.tabs.widget(index)
        if not isinstance(tab, EditorTab):
            return
        if tab.is_dirty and not self._confirm_discard_tab(tab):
            return
        if self.tabs.count() <= 1:
            self._add_new_tab()
        self.tabs.removeTab(index)

    def _on_current_tab_changed(self, index: int) -> None:
        self._update_title()

    def _on_tab_dirty_changed(self, _dirty: bool) -> None:
        tab = self.sender()
        if isinstance(tab, EditorTab):
            idx = self.tabs.indexOf(tab)
            if idx >= 0:
                self.tabs.setTabText(idx, tab.tab_title)
        self._update_title()

    def _on_file_dropped(self, path: str) -> None:
        self._open_path_and_add_tab(path)

    def dragEnterEvent(self, event: QDragEnterEvent) -> None:
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            event.ignore()

    def dropEvent(self, event: QDropEvent) -> None:
        for url in event.mimeData().urls():
            if url.isLocalFile():
                path = url.toLocalFile()
                if path.endswith((".md", ".markdown", ".txt")):
                    self._open_path_and_add_tab(path)
                    event.acceptProposedAction()
                    return
        event.ignore()

    # ----- menus / actions -----

    def _build_menus(self) -> None:
        mb = self.menuBar()

        file_menu = mb.addMenu("文件(&F)")
        self._add_action(file_menu, "新建(&N)", self.new_file, QKeySequence.New)
        self._add_action(file_menu, "打开(&O)…", self.open_file, QKeySequence.Open)
        self._add_action(file_menu, "保存(&S)", self.save_file, QKeySequence.Save)
        self._add_action(file_menu, "另存为(&A)…", self.save_file_as, QKeySequence("Ctrl+Shift+S"))
        file_menu.addSeparator()
        self._recent_menu = file_menu.addMenu("最近文件")
        self._rebuild_recent_menu()
        file_menu.addSeparator()
        self._add_action(file_menu, "退出(&Q)", self.close, QKeySequence("Ctrl+Q"))

        view_menu = mb.addMenu("视图(&V)")
        view_group = QActionGroup(self)
        view_group.setExclusive(True)

        self._translation_action = QAction("翻译对照(&T)", self, checkable=True)
        self._translation_action.setShortcut(QKeySequence("Ctrl+1"))
        self._translation_action.triggered.connect(
            lambda: self._switch_view(VIEW_TRANSLATION)
        )
        view_group.addAction(self._translation_action)
        view_menu.addAction(self._translation_action)

        self._preview_action = QAction("渲染预览(&P)", self, checkable=True)
        self._preview_action.setShortcut(QKeySequence("Ctrl+2"))
        self._preview_action.setChecked(True)
        self._preview_action.triggered.connect(lambda: self._switch_view(VIEW_PREVIEW))
        view_group.addAction(self._preview_action)
        view_menu.addAction(self._preview_action)

        help_menu = mb.addMenu("帮助(&H)")
        self._add_action(help_menu, "关于", self._show_about)

    def _add_action(self, menu, text, slot, shortcut=None) -> QAction:
        act = QAction(text, self)
        if shortcut is not None:
            act.setShortcut(shortcut)
        act.triggered.connect(slot)
        menu.addAction(act)
        return act

    # ----- file ops -----

    def new_file(self) -> None:
        self._add_new_tab()

    def open_file(self, path: str | None = None) -> None:
        if not path:
            path, _ = QFileDialog.getOpenFileName(
                self, "打开 Markdown 文件", "", "Markdown (*.md *.markdown);;所有文件 (*.*)"
            )
        if not path:
            return
        self._open_path_and_add_tab(path)

    def _open_path_and_add_tab(self, path: str) -> None:
        try:
            text = Path(path).read_text(encoding="utf-8")
        except (OSError, UnicodeDecodeError) as e:
            QMessageBox.critical(self, "打开失败", str(e))
            return
        self._add_new_tab(text=text, path=Path(path))

    def save_file(self) -> bool:
        tab = self._current_tab()
        if tab is None:
            return False
        if tab.current_path is None:
            return self.save_file_as()
        return self._write_to(tab, tab.current_path)

    def save_file_as(self) -> bool:
        tab = self._current_tab()
        if tab is None:
            return False
        path, _ = QFileDialog.getSaveFileName(
            self, "另存为", "", "Markdown (*.md);;所有文件 (*.*)"
        )
        if not path:
            return False
        return self._write_to(tab, Path(path))

    def _write_to(self, tab: EditorTab, path: Path) -> bool:
        try:
            path.write_text(tab.markdown_text(), encoding="utf-8")
        except OSError as e:
            QMessageBox.critical(self, "保存失败", str(e))
            return False
        tab.current_path = path
        tab._dirty = False
        tab.dirtyChanged.emit(False)
        idx = self.tabs.indexOf(tab)
        if idx >= 0:
            self.tabs.setTabText(idx, tab.tab_title)
        self._push_recent(str(path))
        self._update_title()
        self.statusBar().showMessage(f"已保存:{path}", 3000)
        return True

    def _confirm_discard_tab(self, tab: EditorTab) -> bool:
        if not tab.is_dirty:
            return True
        name = tab.current_path.name if tab.current_path else "未命名"
        ret = QMessageBox.question(
            self,
            "未保存的更改",
            f"{name} 有未保存的更改,是否保存?",
            QMessageBox.Save | QMessageBox.Discard | QMessageBox.Cancel,
            QMessageBox.Save,
        )
        if ret == QMessageBox.Save:
            return self._write_to(tab, tab.current_path) if tab.current_path else self.save_file_as()
        return ret == QMessageBox.Discard

    def _confirm_discard(self) -> bool:
        for i in range(self.tabs.count()):
            tab = self.tabs.widget(i)
            if isinstance(tab, EditorTab) and tab.is_dirty:
                if not self._confirm_discard_tab(tab):
                    return False
        return True

    # ----- recent files -----

    def _push_recent(self, path: str) -> None:
        items: list[str] = self._settings.value("recent", [], type=list) or []
        if path in items:
            items.remove(path)
        items.insert(0, path)
        items = items[:RECENT_MAX]
        self._settings.setValue("recent", items)
        self._rebuild_recent_menu()

    def _rebuild_recent_menu(self) -> None:
        self._recent_menu.clear()
        items: list[str] = self._settings.value("recent", [], type=list) or []
        if not items:
            act = QAction("(空)", self)
            act.setEnabled(False)
            self._recent_menu.addAction(act)
            return
        for p in items:
            act = QAction(p, self)
            act.triggered.connect(lambda _=False, path=p: self._open_path_and_add_tab(path))
            self._recent_menu.addAction(act)

    # ----- views / rendering -----

    def _switch_view(self, index: int) -> None:
        tab = self._current_tab()
        if tab:
            tab.set_view_index(index)

    # ----- scroll sync -----

    def _sync_scroll(self, _value: int) -> None:
        pass

    # ----- misc -----

    def _update_title(self) -> None:
        tab = self._current_tab()
        name = tab.tab_title if tab else "未命名"
        self.setWindowTitle(f"{name} — {APP_NAME} v{__version__}")

    def _show_about(self) -> None:
        QMessageBox.about(
            self,
            "关于",
            f"<b>{APP_NAME}</b> v{__version__}<br>"
            "轻量 Markdown 编辑器,内置离线英译中。",
        )

    # ----- preview actions -----

    def _open_preview_in_chrome(self) -> None:
        tab = self._current_tab()
        if tab is None:
            return
        html = tab.preview.current_html() or self._renderer.render(tab.markdown_text())
        ok, msg = open_html_in_chrome_translate(html)
        self.statusBar().showMessage(msg, 6000)
        if not ok and "未检测到 Chrome" not in msg:
            QMessageBox.warning(self, "打开失败", msg)

    def _save_preview_as_html(self) -> None:
        tab = self._current_tab()
        if tab is None:
            return
        html = tab.preview.current_html() or self._renderer.render(tab.markdown_text())
        path, _ = QFileDialog.getSaveFileName(
            self, "保存为 HTML", "", "HTML 文件 (*.html);;所有文件 (*.*)"
        )
        if not path:
            return
        try:
            Path(path).write_text(html, encoding="utf-8")
            self.statusBar().showMessage(f"HTML 已保存: {path}", 3000)
        except OSError as e:
            QMessageBox.critical(self, "保存失败", str(e))

    def _save_preview_as_pdf(self) -> None:
        tab = self._current_tab()
        if tab is None:
            return
        path, _ = QFileDialog.getSaveFileName(
            self, "保存为 PDF", "", "PDF 文件 (*.pdf);;所有文件 (*.*)"
        )
        if not path:
            return
        try:
            tab.preview.page().printToPdf(path)
            self.statusBar().showMessage(f"PDF 已保存: {path}", 3000)
        except Exception as e:
            QMessageBox.critical(self, "保存失败", str(e))

    def _save_preview_as_word(self) -> None:
        tab = self._current_tab()
        if tab is None:
            return
        md_text = tab.markdown_text()
        path, _ = QFileDialog.getSaveFileName(
            self, "保存为 Word", "", "Word 文件 (*.docx);;所有文件 (*.*)"
        )
        if not path:
            return
        try:
            doc = Document()
            doc.add_heading("文档标题", 0)
            for line in md_text.splitlines():
                stripped = line.strip()
                if not stripped:
                    doc.add_paragraph("")
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], 3)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], 2)
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], 1)
                else:
                    doc.add_paragraph(stripped)
            doc.save(path)
            self.statusBar().showMessage(f"Word 已保存: {path}", 3000)
        except Exception as e:
            QMessageBox.critical(self, "保存失败", str(e))

    # ----- translation wiring -----

    def _init_worker(self) -> None:
        if not self._dict_path.exists():
            self.statusBar().showMessage(
                "悬停词典未找到 (resources/dict/ecdict.sqlite),悬停翻译已禁用", 8000
            )

        m_path = model_path()
        self._worker: TranslationWorker | None = None
        if m_path.exists():
            self._worker = TranslationWorker(m_path)
            self._worker.translated.connect(self._on_paragraph_translated)
            self._worker.failed.connect(self._on_translation_failed)
            self._worker.start()
            set_active_worker(self._worker)
        else:
            self._translation_action.setEnabled(False)
            self._translation_action.setToolTip(
                "Argos 模型未找到 (resources/models/translate-en_zh*.argosmodel)"
            )
            self._preview_action.setChecked(True)

    def _on_paragraph_translated(self, _original: str, _translation: str) -> None:
        if not self._refresh_pending:
            self._refresh_pending = True
            self._refresh_timer.start(150)

    def _do_refresh_translation(self) -> None:
        self._refresh_pending = False
        tab = self._current_tab()
        if tab and tab.current_view_index() == VIEW_TRANSLATION:
            tab._refresh_right_pane()

    def _on_translation_failed(self, message: str) -> None:
        self.statusBar().showMessage(f"翻译失败: {message}", 6000)

    def closeEvent(self, event) -> None:
        if not self._confirm_discard():
            event.ignore()
            return
        if self._worker is not None:
            set_active_worker(None)
            self._worker.stop()
        event.accept()